# Databricks notebook source
# COMMAND ----------
# Run once if needed, then restart Python from Databricks UI if prompted.
# %pip install -q google-genai beautifulsoup4 lxml nest_asyncio pandas

# COMMAND ----------
import os
import re
import time
import uuid
import random
import asyncio
from datetime import datetime
from typing import List, Tuple, Dict, Any

import pandas as pd
from bs4 import BeautifulSoup, Comment
from pyspark.sql import functions as F

try:
    import nest_asyncio
    nest_asyncio.apply()
except Exception:
    pass

from google import genai
from google.genai import types

# COMMAND ----------
# ---------------------------
# 1) Configuration
# ---------------------------
TABLE_NAME = "default.kb_poc_docs"

ID_COL = "doc_id"          # if missing, notebook auto-generates one
TEXT_COL = "html_text"     # you said your text column name is html_text
CAT_COL = "category"       # optional; defaults to "internal" if missing

N_DOCS = 30                # number of docs to benchmark
MIN_TEXT_CHARS = 2000      # use longer docs to make chunk parallelism visible

# Model/API
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip()
# If env var is not set, paste directly below:
# GEMINI_API_KEY = "YOUR_GEMINI_KEY"

MODEL = "gemini-1.5-flash"   # fast + cheaper for benchmarking
TEMPERATURE = 0.1

# Protected tags: text inside these tags is never rewritten
PROTECTED_TAGS = {"b", "strong"}

# Chunking settings
MAX_SEGMENT_CHARS = 12000
CHUNK_TARGET_CHARS = 3500
CHUNK_HARD_LIMIT_CHARS = 5000
HARMONIZE_GROUP_CHARS = 8000

# Retry settings
MAX_RETRIES = 4
RETRY_BASE_SECONDS = 1.5

# Parallel benchmark variants
EXPERIMENTS = [
    {"variant_id": "A_seq", "max_doc_concurrency": 1, "max_request_concurrency": 1},
    {"variant_id": "B_chunk_parallel", "max_doc_concurrency": 1, "max_request_concurrency": 4},
    {"variant_id": "C_doc_and_chunk_parallel", "max_doc_concurrency": 3, "max_request_concurrency": 8},
]

assert GEMINI_API_KEY, "Set GEMINI_API_KEY first."
client = genai.Client(api_key=GEMINI_API_KEY)

# COMMAND ----------
# ---------------------------
# 2) Load sample docs
# ---------------------------
df = spark.table(TABLE_NAME)
cols = set(df.columns)

if TEXT_COL not in cols:
    raise ValueError(f"Column '{TEXT_COL}' not found in table {TABLE_NAME}. Available columns: {sorted(cols)}")

select_cols = []
if ID_COL in cols:
    select_cols.append(F.col(ID_COL).cast("string").alias("doc_id"))
else:
    select_cols.append(F.monotonically_increasing_id().cast("string").alias("doc_id"))

select_cols.append(F.col(TEXT_COL).cast("string").alias("html_text"))

if CAT_COL in cols:
    select_cols.append(F.col(CAT_COL).cast("string").alias("category"))
else:
    select_cols.append(F.lit("internal").alias("category"))

sample_df = (
    df.select(*select_cols)
      .filter(F.col("html_text").isNotNull())
      .withColumn("char_len", F.length("html_text"))
      .filter(F.col("char_len") >= MIN_TEXT_CHARS)
      .orderBy(F.rand(42))
      .limit(N_DOCS)
)

rows = sample_df.collect()
docs: List[Tuple[str, str, str]] = [
    (r["doc_id"], r["html_text"], r["category"] or "internal")
    for r in rows
]

if not docs:
    raise ValueError("No docs matched MIN_TEXT_CHARS filter. Lower MIN_TEXT_CHARS and rerun.")

print(f"Loaded {len(docs)} docs for benchmark.")
display(sample_df.select("doc_id", "category", "char_len"))

# COMMAND ----------
# ---------------------------
# 3) HTML + chunking helpers
# ---------------------------
class LLMCallError(Exception):
    pass

_SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?])\s+")

def gemini_text(resp) -> str:
    txt = getattr(resp, "text", None)
    if txt:
        return txt
    try:
        parts = []
        for cand in (resp.candidates or []):
            content = getattr(cand, "content", None)
            if not content:
                continue
            for part in (getattr(content, "parts", None) or []):
                t = getattr(part, "text", None)
                if t:
                    parts.append(t)
        return "\n".join(parts)
    except Exception:
        return ""

def is_protected_node(node) -> bool:
    p = node.parent
    while p is not None and getattr(p, "name", None):
        if p.name and p.name.lower() in PROTECTED_TAGS:
            return True
        p = p.parent
    return False

def parse_segments(html_text: str):
    # Parses HTML and returns all non-empty text nodes as editable segments.
    soup = BeautifulSoup(html_text, "lxml")
    segments = []
    idx = 0
    for node in soup.find_all(string=True):
        if isinstance(node, Comment):
            continue
        txt = str(node)
        if not txt.strip():
            continue
        segments.append({
            "seg_id": idx,
            "node": node,
            "text": txt,
            "is_protected": is_protected_node(node),
        })
        idx += 1
    return soup, segments

def force_split_text(text: str, hard_limit: int) -> List[str]:
    parts = []
    n = len(text)
    i = 0
    while i < n:
        end = min(i + hard_limit, n)
        if end < n:
            cut = text.rfind(" ", i, end)
            if cut > i + hard_limit // 2:
                end = cut
        part = text[i:end].strip()
        if part:
            parts.append(part)
        i = end
        while i < n and text[i].isspace():
            i += 1
    return parts

def split_large_paragraph(paragraph: str, hard_limit: int) -> List[str]:
    sentences = _SENTENCE_SPLIT_RE.split(paragraph.strip())
    out = []
    current = ""
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        if len(s) > hard_limit:
            if current:
                out.append(current)
                current = ""
            out.extend(force_split_text(s, hard_limit))
            continue
        candidate = s if not current else f"{current} {s}"
        if len(candidate) <= hard_limit:
            current = candidate
        else:
            out.append(current)
            current = s
    if current:
        out.append(current)
    return out

def split_text_for_rewrite(text: str, target_chars: int, hard_limit: int) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []

    # Split by paragraph blocks first
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]

    atomic_parts = []
    for p in paragraphs:
        if len(p) <= hard_limit:
            atomic_parts.append(p)
        else:
            atomic_parts.extend(split_large_paragraph(p, hard_limit))

    # Pack to near target size
    chunks = []
    current = ""
    for part in atomic_parts:
        sep = "\n\n" if current else ""
        candidate = f"{current}{sep}{part}"
        if len(candidate) <= target_chars:
            current = candidate
        else:
            if current:
                chunks.append(current)
                current = part
            else:
                chunks.append(part)
                current = ""
    if current:
        chunks.append(current)

    # Final hard safety
    safe = []
    for c in chunks:
        if len(c) <= hard_limit:
            safe.append(c)
        else:
            safe.extend(force_split_text(c, hard_limit))
    return safe

def pack_text_groups(texts: List[str], max_chars: int) -> List[List[str]]:
    groups = []
    current = []
    current_len = 0
    for t in texts:
        add_len = len(t) + (2 if current else 0)
        if current and (current_len + add_len > max_chars):
            groups.append(current)
            current = [t]
            current_len = len(t)
        else:
            current.append(t)
            current_len += add_len
    if current:
        groups.append(current)
    return groups

async def gather_in_batches(coros, batch_size: int = 200):
    results = []
    for i in range(0, len(coros), batch_size):
        results.extend(await asyncio.gather(*coros[i:i + batch_size]))
    return results

def normalize_ws(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()

def protected_texts(html_text: str) -> List[str]:
    soup = BeautifulSoup(html_text, "lxml")
    out = []
    for t in soup.find_all(PROTECTED_TAGS):
        for n in t.find_all(string=True):
            v = normalize_ws(str(n))
            if v:
                out.append(v)
    return out

# COMMAND ----------
# ---------------------------
# 4) Async LLM calls
# ---------------------------
async def generate_with_retry_async(
    user_prompt: str,
    system_instruction: str,
    semaphore: asyncio.Semaphore,
    max_output_tokens: int = 4096,
) -> str:
    last_err = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            cfg = types.GenerateContentConfig(
                system_instruction=system_instruction,
                temperature=TEMPERATURE,
                max_output_tokens=max_output_tokens,
            )
            async with semaphore:
                resp = await asyncio.to_thread(
                    client.models.generate_content,
                    model=MODEL,
                    contents=user_prompt,
                    config=cfg,
                )
            return (gemini_text(resp) or "").strip()
        except Exception as e:
            last_err = e
            if attempt == MAX_RETRIES:
                break
            sleep_s = RETRY_BASE_SECONDS * (2 ** (attempt - 1)) + random.random()
            await asyncio.sleep(sleep_s)

    raise LLMCallError(str(last_err)) from last_err

async def rewrite_chunk_async(
    chunk_text: str,
    category: str,
    chunk_idx: int,
    chunk_total: int,
    semaphore: asyncio.Semaphore,
) -> str:
    system_instruction = (
        "You are an enterprise technical editor. Rewrite for clarity and professionalism. "
        "Do not change meaning or facts. Do not change numbers, dates, monetary amounts, URLs, IDs, or deadlines. "
        "Return plain text only."
    )
    user_prompt = f"""
Category: {category}
Chunk: {chunk_idx}/{chunk_total}

Rewrite this text:
{chunk_text}
""".strip()

    out = await generate_with_retry_async(user_prompt, system_instruction, semaphore)
    return out if out else chunk_text

async def harmonize_group_async(
    group_texts: List[str],
    category: str,
    semaphore: asyncio.Semaphore,
) -> str:
    joined = "\n\n".join(group_texts)
    system_instruction = (
        "You are an enterprise editor performing harmonization. "
        "Merge rewritten chunks into one coherent text with consistent tone and flow. "
        "Do not change meaning, facts, numbers, dates, URLs, IDs, or deadlines. "
        "Do not add new content. Return plain text only."
    )
    user_prompt = f"""
Category: {category}

Harmonize these rewritten chunks into one continuous text:
{joined}
""".strip()

    try:
        out = await generate_with_retry_async(user_prompt, system_instruction, semaphore)
        return out if out else joined
    except Exception:
        # Fail-open to already rewritten text
        return joined

# COMMAND ----------
# ---------------------------
# 5) Rewriter (segment -> html -> batch)
# ---------------------------
async def rewrite_segment_async(
    text: str,
    category: str,
    semaphore: asyncio.Semaphore,
) -> str:
    raw = text or ""
    if not raw.strip():
        return raw

    # Preserve surrounding whitespace
    leading_ws = raw[: len(raw) - len(raw.lstrip())]
    trailing_ws = raw[len(raw.rstrip()):]
    core = raw.strip()

    # Small segment => single request
    if len(core) <= MAX_SEGMENT_CHARS:
        rewritten = await rewrite_chunk_async(core, category, 1, 1, semaphore)
        return f"{leading_ws}{rewritten}{trailing_ws}"

    # Large segment => chunk + parallel rewrite + hierarchical harmonize
    chunks = split_text_for_rewrite(core, CHUNK_TARGET_CHARS, CHUNK_HARD_LIMIT_CHARS)
    if not chunks:
        return raw

    coros = [
        rewrite_chunk_async(ch, category, i + 1, len(chunks), semaphore)
        for i, ch in enumerate(chunks)
    ]
    rewritten_chunks = await gather_in_batches(coros, batch_size=100)

    current = rewritten_chunks
    while len(current) > 1:
        groups = pack_text_groups(current, HARMONIZE_GROUP_CHARS)
        coros = [harmonize_group_async(g, category, semaphore) for g in groups]
        current = await gather_in_batches(coros, batch_size=50)

    final_core = current[0] if current else core
    return f"{leading_ws}{final_core}{trailing_ws}"

async def rewrite_html_async(
    html_text: str,
    category: str,
    request_semaphore: asyncio.Semaphore,
) -> Tuple[str, Dict[str, Any]]:
    soup, segments = parse_segments(html_text)

    editable_idxs = []
    coros = []
    protected_count = 0

    for i, seg in enumerate(segments):
        if seg["is_protected"]:
            protected_count += 1
            continue
        editable_idxs.append(i)
        coros.append(rewrite_segment_async(seg["text"], category, request_semaphore))

    rewritten_texts = []
    if coros:
        rewritten_texts = await gather_in_batches(coros, batch_size=200)

    for i, new_text in zip(editable_idxs, rewritten_texts):
        segments[i]["node"].replace_with(new_text)

    stats = {
        "total_segments": len(segments),
        "editable_segments": len(editable_idxs),
        "protected_segments": protected_count,
    }
    return str(soup), stats

async def rewrite_document_task(
    doc_id: str,
    html_text: str,
    category: str,
    request_semaphore: asyncio.Semaphore,
) -> Dict[str, Any]:
    t0 = time.perf_counter()
    original_protected = protected_texts(html_text)

    try:
        rewritten_html, seg_stats = await rewrite_html_async(html_text, category, request_semaphore)
        status = "ok"
        error = ""
    except Exception as e:
        rewritten_html = html_text
        seg_stats = {"total_segments": 0, "editable_segments": 0, "protected_segments": 0}
        status = "error"
        error = str(e)

    t1 = time.perf_counter()
    rewritten_protected = protected_texts(rewritten_html)

    return {
        "doc_id": doc_id,
        "category": category,
        "status": status,
        "error": error,
        "elapsed_sec": round(t1 - t0, 4),
        "input_chars": len(html_text or ""),
        "output_chars": len(rewritten_html or ""),
        "changed": rewritten_html != html_text,
        "protected_unchanged": original_protected == rewritten_protected,
        "total_segments": seg_stats["total_segments"],
        "editable_segments": seg_stats["editable_segments"],
        "protected_segments": seg_stats["protected_segments"],
        "rewritten_html": rewritten_html,
    }

async def rewrite_html_batch_async(
    docs: List[Tuple[str, str, str]],  # (doc_id, html_text, category)
    max_doc_concurrency: int,
    max_request_concurrency: int,
) -> List[Dict[str, Any]]:
    request_sem = asyncio.Semaphore(max_request_concurrency)
    doc_sem = asyncio.Semaphore(max_doc_concurrency)

    async def run_one(doc):
        doc_id, html_text, category = doc
        async with doc_sem:
            return await rewrite_document_task(
                doc_id=doc_id,
                html_text=html_text,
                category=category,
                request_semaphore=request_sem,
            )

    coros = [run_one(d) for d in docs]
    return await gather_in_batches(coros, batch_size=max_doc_concurrency * 10)

# COMMAND ----------
# ---------------------------
# 6) Benchmark runner
# ---------------------------
async def run_experiment(
    docs: List[Tuple[str, str, str]],
    variant_id: str,
    max_doc_concurrency: int,
    max_request_concurrency: int,
) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    start = time.perf_counter()
    doc_rows = await rewrite_html_batch_async(
        docs,
        max_doc_concurrency=max_doc_concurrency,
        max_request_concurrency=max_request_concurrency,
    )
    wall = time.perf_counter() - start

    total = len(doc_rows)
    ok = sum(1 for r in doc_rows if r["status"] == "ok")
    err = total - ok
    changed = sum(1 for r in doc_rows if r["changed"])
    protected_ok = sum(1 for r in doc_rows if r["protected_unchanged"])
    total_chars = sum(r["input_chars"] for r in doc_rows)
    avg_doc_sec = sum(r["elapsed_sec"] for r in doc_rows) / max(total, 1)

    summary = {
        "variant_id": variant_id,
        "max_doc_concurrency": max_doc_concurrency,
        "max_request_concurrency": max_request_concurrency,
        "docs_total": total,
        "docs_ok": ok,
        "docs_error": err,
        "wall_time_sec": round(wall, 4),
        "avg_doc_elapsed_sec": round(avg_doc_sec, 4),
        "docs_per_min": round((ok / wall) * 60, 4) if wall > 0 else 0.0,
        "chars_per_sec": round(total_chars / wall, 4) if wall > 0 else 0.0,
        "changed_rate": round(changed / max(total, 1), 4),
        "protected_unchanged_rate": round(protected_ok / max(total, 1), 4),
    }

    for r in doc_rows:
        r["variant_id"] = variant_id

    return doc_rows, summary

async def run_all_experiments(docs, experiments):
    all_doc_rows = []
    summary_rows = []
    for exp in experiments:
        print(f"Running {exp['variant_id']} ...")
        doc_rows, summary = await run_experiment(
            docs=docs,
            variant_id=exp["variant_id"],
            max_doc_concurrency=exp["max_doc_concurrency"],
            max_request_concurrency=exp["max_request_concurrency"],
        )
        all_doc_rows.extend(doc_rows)
        summary_rows.append(summary)
    return all_doc_rows, summary_rows

# COMMAND ----------
# ---------------------------
# 7) Execute benchmark
# ---------------------------
def run_async(coro):
    try:
        loop = asyncio.get_event_loop()
        return loop.run_until_complete(coro)
    except RuntimeError:
        return asyncio.run(coro)

all_doc_rows, summary_rows = run_async(run_all_experiments(docs, EXPERIMENTS))

summary_pdf = pd.DataFrame(summary_rows).sort_values("docs_per_min", ascending=False)
display(summary_pdf)

# COMMAND ----------
# ---------------------------
# 8) Save results to Delta tables
# ---------------------------
RUN_ID = f"parallel_test_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
RUN_TS_UTC = datetime.utcnow().isoformat()

for r in summary_rows:
    r["run_id"] = RUN_ID
    r["run_ts_utc"] = RUN_TS_UTC

for r in all_doc_rows:
    r["run_id"] = RUN_ID
    r["run_ts_utc"] = RUN_TS_UTC

summary_table = "default.kb_parallel_benchmark_summary"
detail_table = "default.kb_parallel_benchmark_detail"

spark.createDataFrame(summary_rows).write.mode("append").saveAsTable(summary_table)
spark.createDataFrame(all_doc_rows).write.mode("append").saveAsTable(detail_table)

print("Saved summary table:", summary_table)
print("Saved detail table :", detail_table)
print("RUN_ID:", RUN_ID)

# COMMAND ----------
# ---------------------------
# 9) Inspect results
# ---------------------------
display(
    spark.table(summary_table)
         .filter(F.col("run_id") == RUN_ID)
         .orderBy(F.col("docs_per_min").desc())
)

display(
    spark.table(detail_table)
         .filter(F.col("run_id") == RUN_ID)
         .select(
             "variant_id", "doc_id", "category", "status", "elapsed_sec",
             "input_chars", "output_chars", "changed", "protected_unchanged"
         )
         .orderBy("variant_id", "elapsed_sec")
)

# COMMAND ----------
# ---------------------------
# 10) Save rewritten HTML for best variant
# ---------------------------
best_variant = summary_pdf.iloc[0]["variant_id"]
best_rows = [r for r in all_doc_rows if r["variant_id"] == best_variant and r["status"] == "ok"]

best_out_table = "default.kb_poc_rewritten_best_variant"
spark.createDataFrame(best_rows).select(
    "doc_id", "category", "variant_id", "rewritten_html", "run_id", "run_ts_utc"
).write.mode("overwrite").saveAsTable(best_out_table)

print("Best variant:", best_variant)
print("Saved rewritten outputs to:", best_out_table)






# ===============================================================
# DOCX Transformation with OpenAI + LangChain chunking + harmonization
# Beginner-friendly version with detailed comments
# ===============================================================
# Install once if needed:
# %pip install -q openai python-docx nest_asyncio langchain-text-splitters

import os
import re
import random
import asyncio
from typing import List

from docx import Document
from openai import AsyncOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter

# If you are in a notebook, this avoids "event loop already running" errors.
try:
    import nest_asyncio
    nest_asyncio.apply()
except Exception:
    pass


# ===============================================================
# 1) USER CONFIGURATION
# ===============================================================

# Input/output Word file paths
pathStr = "/dbfs/FileStore/input/my_document.docx"
output_path = "/dbfs/FileStore/output/my_document_rewritten.docx"

# Rules file: one rule per line (supports 100+ rules)
rules_file_path = "/dbfs/FileStore/input/rules.txt"

# Document category used in prompt context
document_category = "internal"  # internal / process_instructions / customer_qa

# OpenAI settings
# Set key first: os.environ["OPENAI_API_KEY"] = "sk-..."
MODEL_NAME = "gpt-4.1-mini"
TEMPERATURE = 0.1
MAX_OUTPUT_TOKENS = 2500

# Concurrency settings
MAX_CONCURRENT_REQUESTS = 6      # max in-flight OpenAI calls at once
ASYNC_BATCH_SIZE = 120           # how many tasks to await per batch

# Retry settings
MAX_RETRIES = 4
RETRY_BASE_SECONDS = 1.5

# Chunking settings
MAX_TEXT_WITHOUT_CHUNKING = 8000  # if text <= this, one call (no chunking)
TARGET_CHUNK_SIZE = 2500          # preferred chunk size for splitter
CHUNK_OVERLAP = 0                 # for rewrite pipelines, keep this 0
HARD_CHUNK_LIMIT = 3500           # final hard safety per chunk

# Harmonization settings
ENABLE_HARMONIZATION = True
MAX_HARMONIZE_GROUP_CHARS = 10000

# Formatting protection
PROTECT_BOLD_TEXT = True          # do not modify run.bold == True

# Optional overlap de-dup guard while combining chunks
MAX_OVERLAP_DEDUP_CHARS = 300


# ===============================================================
# 2) RULES HELPERS
# ===============================================================

def load_rules_from_file(file_path: str) -> List[str]:
    """
    Load rules from a plain text file.
    - One rule per line
    - Empty lines ignored
    - Lines starting with # are comments and ignored
    """
    rules = []
    with open(file_path, "r", encoding="utf-8") as file_obj:
        for line in file_obj:
            cleaned = line.strip()
            if cleaned == "":
                continue
            if cleaned.startswith("#"):
                continue
            rules.append(cleaned)
    return rules


def rules_to_numbered_text(rules_list: List[str]) -> str:
    """
    Convert list of rules into a numbered block.
    This is easier for the model to follow.
    """
    lines = []
    index = 1
    for rule in rules_list:
        lines.append(str(index) + ". " + rule)
        index = index + 1
    return "\n".join(lines)


# ===============================================================
# 3) CHUNKING HELPERS (LangChain splitter + hard safety)
# ===============================================================

def create_langchain_splitter() -> RecursiveCharacterTextSplitter:
    """
    Create a recursive splitter.
    It tries larger boundaries first (paragraph, newline), then smaller.
    """
    splitter = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", ". ", " ", ""],
        chunk_size=TARGET_CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len
    )
    return splitter


def hard_split_text(text_value: str, max_chars: int) -> List[str]:
    """
    Last-resort split by character count.
    Used only if a chunk still exceeds HARD_CHUNK_LIMIT.
    """
    pieces = []
    start = 0
    text_value = text_value or ""
    while start < len(text_value):
        end = start + max_chars
        piece = text_value[start:end].strip()
        if piece != "":
            pieces.append(piece)
        start = end
    return pieces


def split_text_into_chunks_langchain(text_value: str, splitter: RecursiveCharacterTextSplitter) -> List[str]:
    """
    Split text with LangChain and enforce hard size limit.
    """
    text_value = (text_value or "").strip()
    if text_value == "":
        return []

    raw_chunks = splitter.split_text(text_value)
    safe_chunks = []

    for chunk in raw_chunks:
        chunk = (chunk or "").strip()
        if chunk == "":
            continue

        if len(chunk) <= HARD_CHUNK_LIMIT:
            safe_chunks.append(chunk)
        else:
            # If a chunk is still too big, force split it
            forced = hard_split_text(chunk, HARD_CHUNK_LIMIT)
            for piece in forced:
                safe_chunks.append(piece)

    return safe_chunks


# ===============================================================
# 4) PROMPT BUILDERS
# ===============================================================

def build_rewrite_user_message(category: str, rules_text: str, chunk_text: str, chunk_number: int, total_chunks: int) -> str:
    """
    Build an easy-to-read rewrite prompt.
    """
    lines = []
    lines.append("Document category: " + str(category))
    lines.append("Chunk number: " + str(chunk_number) + " of " + str(total_chunks))
    lines.append("")
    lines.append("Apply ALL rules below while preserving meaning and facts.")
    lines.append("Rules:")
    lines.append(rules_text)
    lines.append("")
    lines.append("Text to rewrite:")
    lines.append(chunk_text)
    return "\n".join(lines)


def build_harmonize_user_message(category: str, rules_text: str, merged_group_text: str, level_number: int, group_number: int, total_groups: int) -> str:
    """
    Build harmonization prompt for one group.
    """
    lines = []
    lines.append("Document category: " + str(category))
    lines.append("Harmonization level: " + str(level_number))
    lines.append("Group number: " + str(group_number) + " of " + str(total_groups))
    lines.append("")
    lines.append("Task: Make tone and flow consistent across this text.")
    lines.append("Do NOT change meaning, facts, numbers, dates, URLs, IDs, or deadlines.")
    lines.append("Do NOT add new information.")
    lines.append("")
    lines.append("Rules to respect:")
    lines.append(rules_text)
    lines.append("")
    lines.append("Text to harmonize:")
    lines.append(merged_group_text)
    return "\n".join(lines)


# ===============================================================
# 5) OPENAI CALL HELPERS
# ===============================================================

def extract_response_text(response) -> str:
    """
    Safely extract text from OpenAI Responses API output.
    """
    direct = getattr(response, "output_text", None)
    if direct:
        return direct

    parts = []
    output_items = getattr(response, "output", None) or []
    for item in output_items:
        content = getattr(item, "content", None) or []
        for c in content:
            text_val = getattr(c, "text", None)
            if text_val:
                parts.append(text_val)

    return "\n".join(parts).strip()


async def call_openai_with_retry(
    client: AsyncOpenAI,
    semaphore: asyncio.Semaphore,
    system_message: str,
    user_message: str
) -> str:
    """
    Make one OpenAI call with:
    - Semaphore control (parallelism cap)
    - Retry/backoff for temporary failures
    """
    last_error = None
    attempt = 1

    while attempt <= MAX_RETRIES:
        try:
            async with semaphore:
                response = await client.responses.create(
                    model=MODEL_NAME,
                    temperature=TEMPERATURE,
                    max_output_tokens=MAX_OUTPUT_TOKENS,
                    input=[
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": user_message},
                    ],
                )

            text_out = extract_response_text(response).strip()
            return text_out

        except Exception as err:
            last_error = err

            if attempt == MAX_RETRIES:
                break

            sleep_seconds = RETRY_BASE_SECONDS * (2 ** (attempt - 1)) + random.random()
            await asyncio.sleep(sleep_seconds)
            attempt = attempt + 1

    raise RuntimeError("OpenAI call failed after retries: " + str(last_error))


async def gather_in_batches(coros: List[asyncio.Future], batch_size: int, return_exceptions: bool = True):
    """
    Await tasks in smaller groups to avoid very large gather calls.
    """
    all_results = []
    start = 0

    while start < len(coros):
        end = start + batch_size
        batch = coros[start:end]
        batch_results = await asyncio.gather(*batch, return_exceptions=return_exceptions)
        all_results.extend(batch_results)
        start = end

    return all_results


# ===============================================================
# 6) REWRITE + HARMONIZATION HELPERS
# ===============================================================

def keep_original_outer_spaces(original_text: str, rewritten_core: str) -> str:
    """
    Preserve leading/trailing whitespace from original text.
    """
    if original_text is None:
        return rewritten_core or ""

    left_spaces = original_text[:len(original_text) - len(original_text.lstrip())]
    right_spaces = original_text[len(original_text.rstrip()):]
    middle = (rewritten_core or "").strip()
    return left_spaces + middle + right_spaces


def longest_suffix_prefix_overlap(left_text: str, right_text: str, max_check: int) -> int:
    """
    Find overlap size where suffix(left_text) == prefix(right_text).
    Helps avoid accidental duplication if overlap is used.
    """
    max_possible = min(len(left_text), len(right_text), max_check)

    # Require at least small overlap size to avoid random small matches.
    minimum_overlap = 20
    if max_possible < minimum_overlap:
        return 0

    size = max_possible
    while size >= minimum_overlap:
        if left_text[-size:] == right_text[:size]:
            return size
        size = size - 1

    return 0


def combine_rewritten_chunks(rewritten_chunks: List[str]) -> str:
    """
    Combine chunk outputs into one text.
    - remove empty chunks
    - join in order
    - optional overlap de-dup
    - normalize too many blank lines
    """
    cleaned = []
    for chunk in rewritten_chunks:
        value = (chunk or "").strip()
        if value != "":
            cleaned.append(value)

    if len(cleaned) == 0:
        return ""
    if len(cleaned) == 1:
        return cleaned[0]

    merged = cleaned[0]
    index = 1

    while index < len(cleaned):
        next_chunk = cleaned[index]

        overlap_size = longest_suffix_prefix_overlap(
            merged,
            next_chunk,
            MAX_OVERLAP_DEDUP_CHARS
        )

        if overlap_size > 0:
            merged = merged + next_chunk[overlap_size:]
        else:
            merged = merged + "\n\n" + next_chunk

        index = index + 1

    merged = re.sub(r"\n{3,}", "\n\n", merged).strip()
    return merged


async def rewrite_one_chunk(
    client: AsyncOpenAI,
    semaphore: asyncio.Semaphore,
    category: str,
    rules_text: str,
    chunk_text: str,
    chunk_number: int,
    total_chunks: int
) -> str:
    """
    Rewrite one chunk.
    """
    system_message = (
        "You are an enterprise editor.\n"
        "Rewrite for clarity and professionalism.\n"
        "Do not change meaning or facts.\n"
        "Do not change numbers, dates, monetary amounts, URLs, IDs, or deadlines.\n"
        "Return plain text only."
    )

    user_message = build_rewrite_user_message(
        category=category,
        rules_text=rules_text,
        chunk_text=chunk_text,
        chunk_number=chunk_number,
        total_chunks=total_chunks
    )

    rewritten = await call_openai_with_retry(client, semaphore, system_message, user_message)
    if rewritten.strip() == "":
        return chunk_text
    return rewritten


def create_harmonize_groups(text_blocks: List[str], max_chars: int) -> List[List[str]]:
    """
    Group text blocks so each group is model-safe in size.
    """
    groups = []
    current_group = []
    current_len = 0

    for block in text_blocks:
        block_len = len(block)
        separator_len = 2 if len(current_group) > 0 else 0
        candidate_len = current_len + separator_len + block_len

        if candidate_len <= max_chars:
            current_group.append(block)
            current_len = candidate_len
        else:
            if len(current_group) > 0:
                groups.append(current_group)
            current_group = [block]
            current_len = block_len

    if len(current_group) > 0:
        groups.append(current_group)

    return groups


async def harmonize_one_group(
    client: AsyncOpenAI,
    semaphore: asyncio.Semaphore,
    category: str,
    rules_text: str,
    group_texts: List[str],
    level_number: int,
    group_number: int,
    total_groups: int
) -> str:
    """
    Harmonize one group of already-rewritten chunks.
    """
    merged_group_text = combine_rewritten_chunks(group_texts)

    system_message = (
        "You are an enterprise editor.\n"
        "Harmonize style and flow only.\n"
        "Do not change meaning or facts.\n"
        "Do not change numbers, dates, monetary amounts, URLs, IDs, or deadlines.\n"
        "Do not add new information.\n"
        "Return plain text only."
    )

    user_message = build_harmonize_user_message(
        category=category,
        rules_text=rules_text,
        merged_group_text=merged_group_text,
        level_number=level_number,
        group_number=group_number,
        total_groups=total_groups
    )

    try:
        output_text = await call_openai_with_retry(client, semaphore, system_message, user_message)
        if output_text.strip() == "":
            return merged_group_text
        return output_text
    except Exception:
        # Fail-open: if harmonization fails, keep original group text
        return merged_group_text


async def harmonize_chunks_hierarchically(
    client: AsyncOpenAI,
    semaphore: asyncio.Semaphore,
    category: str,
    rules_text: str,
    rewritten_chunks: List[str]
) -> str:
    """
    Hierarchical harmonization:
    - harmonize groups in parallel
    - then harmonize the group outputs
    - continue until one final text remains
    """
    if len(rewritten_chunks) == 0:
        return ""
    if len(rewritten_chunks) == 1:
        return rewritten_chunks[0]

    current_blocks = rewritten_chunks
    level_number = 1

    while len(current_blocks) > 1:
        groups = create_harmonize_groups(current_blocks, MAX_HARMONIZE_GROUP_CHARS)

        tasks = []
        total_groups = len(groups)
        group_index = 1

        for group in groups:
            task = harmonize_one_group(
                client=client,
                semaphore=semaphore,
                category=category,
                rules_text=rules_text,
                group_texts=group,
                level_number=level_number,
                group_number=group_index,
                total_groups=total_groups
            )
            tasks.append(task)
            group_index = group_index + 1

        results = await gather_in_batches(tasks, ASYNC_BATCH_SIZE, return_exceptions=True)

        next_blocks = []
        index = 0
        while index < len(results):
            result_item = results[index]
            original_group_text = combine_rewritten_chunks(groups[index])

            if isinstance(result_item, Exception):
                next_blocks.append(original_group_text)
            else:
                value = (result_item or "").strip()
                if value == "":
                    next_blocks.append(original_group_text)
                else:
                    next_blocks.append(value)

            index = index + 1

        current_blocks = next_blocks
        level_number = level_number + 1

    return current_blocks[0]


async def rewrite_text_with_rules(
    client: AsyncOpenAI,
    semaphore: asyncio.Semaphore,
    splitter: RecursiveCharacterTextSplitter,
    original_text: str,
    category: str,
    rules_text: str
) -> str:
    """
    Main rewrite logic for one text unit.
    - short text => one call
    - long text => chunk + parallel rewrite + optional harmonization
    """
    raw = original_text or ""
    core = raw.strip()

    if core == "":
        return raw

    # Short case: one API call
    if len(core) <= MAX_TEXT_WITHOUT_CHUNKING:
        rewritten = await rewrite_one_chunk(
            client=client,
            semaphore=semaphore,
            category=category,
            rules_text=rules_text,
            chunk_text=core,
            chunk_number=1,
            total_chunks=1
        )
        return keep_original_outer_spaces(raw, rewritten)

    # Long case: split and rewrite chunks in parallel
    chunks = split_text_into_chunks_langchain(core, splitter)
    if len(chunks) == 0:
        return raw

    tasks = []
    total_chunks = len(chunks)
    chunk_index = 1

    for chunk in chunks:
        task = rewrite_one_chunk(
            client=client,
            semaphore=semaphore,
            category=category,
            rules_text=rules_text,
            chunk_text=chunk,
            chunk_number=chunk_index,
            total_chunks=total_chunks
        )
        tasks.append(task)
        chunk_index = chunk_index + 1

    results = await gather_in_batches(tasks, ASYNC_BATCH_SIZE, return_exceptions=True)

    # Fail-open: keep original chunk if that chunk rewrite failed
    rewritten_chunks = []
    i = 0
    while i < len(results):
        result_item = results[i]
        original_chunk = chunks[i]

        if isinstance(result_item, Exception):
            rewritten_chunks.append(original_chunk)
        else:
            out = (result_item or "").strip()
            if out == "":
                rewritten_chunks.append(original_chunk)
            else:
                rewritten_chunks.append(out)
        i = i + 1

    merged = combine_rewritten_chunks(rewritten_chunks)

    # Optional harmonization pass
    if ENABLE_HARMONIZATION and len(rewritten_chunks) > 1:
        harmonized = await harmonize_chunks_hierarchically(
            client=client,
            semaphore=semaphore,
            category=category,
            rules_text=rules_text,
            rewritten_chunks=rewritten_chunks
        )
        if (harmonized or "").strip() != "":
            merged = harmonized.strip()

    return keep_original_outer_spaces(raw, merged)


# ===============================================================
# 7) DOCX TRAVERSAL HELPERS
# ===============================================================

def iter_table_paragraphs(table_obj):
    """
    Yield paragraphs from a table recursively.
    """
    for row in table_obj.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested_table in cell.tables:
                for nested_p in iter_table_paragraphs(nested_table):
                    yield nested_p


def iter_all_paragraphs(doc_obj: Document):
    """
    Yield paragraphs from:
    - main document body
    - body tables
    - headers and footers (and their tables)
    """
    # Body paragraphs
    for p in doc_obj.paragraphs:
        yield p

    # Body tables
    for table in doc_obj.tables:
        for p in iter_table_paragraphs(table):
            yield p

    # Headers/Footers
    for section in doc_obj.sections:
        for p in section.header.paragraphs:
            yield p
        for t in section.header.tables:
            for p in iter_table_paragraphs(t):
                yield p

        for p in section.footer.paragraphs:
            yield p
        for t in section.footer.tables:
            for p in iter_table_paragraphs(t):
                yield p


# ===============================================================
# 8) MAIN TRANSFORMATION FUNCTION
# ===============================================================

async def transform_word_document(
    input_path: str,
    output_file_path: str,
    category: str,
    rules_path: str
):
    """
    End-to-end transform:
    1) Load rules
    2) Load DOCX
    3) Collect editable runs (skip protected bold runs)
    4) Rewrite in parallel with semaphore control
    5) Save rewritten DOCX
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY is not set.")

    # Setup clients/helpers once
    client = AsyncOpenAI(api_key=api_key)
    semaphore = asyncio.Semaphore(MAX_CONCURRENT_REQUESTS)
    splitter = create_langchain_splitter()

    # Load rules
    rules_list = load_rules_from_file(rules_path)
    rules_text = rules_to_numbered_text(rules_list)

    # Load document
    doc = Document(input_path)

    # Collect rewrite jobs
    # Each job corresponds to one editable run in the DOCX
    jobs = []
    total_runs_scanned = 0
    protected_runs = 0

    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            total_runs_scanned = total_runs_scanned + 1
            run_text = run.text or ""

            if run_text.strip() == "":
                continue

            # Protect bold runs exactly as requested
            if PROTECT_BOLD_TEXT and (run.bold is True):
                protected_runs = protected_runs + 1
                continue

            jobs.append({
                "run_obj": run,
                "original_text": run_text
            })

    # Build async tasks
    tasks = []
    for job in jobs:
        task = rewrite_text_with_rules(
            client=client,
            semaphore=semaphore,
            splitter=splitter,
            original_text=job["original_text"],
            category=category,
            rules_text=rules_text
        )
        tasks.append(task)

    # Execute tasks in parallel batches
    results = await gather_in_batches(tasks, ASYNC_BATCH_SIZE, return_exceptions=True)

    # Write results back into runs
    rewritten_runs = 0
    failed_runs = 0

    index = 0
    while index < len(jobs):
        run_obj = jobs[index]["run_obj"]
        original_text = jobs[index]["original_text"]
        result_item = results[index]

        if isinstance(result_item, Exception):
            run_obj.text = original_text  # fail-open
            failed_runs = failed_runs + 1
        else:
            new_text = result_item
            run_obj.text = new_text
            if new_text != original_text:
                rewritten_runs = rewritten_runs + 1

        index = index + 1

    # Save output DOCX
    doc.save(output_file_path)

    # Return summary stats
    stats = {
        "input_path": input_path,
        "output_path": output_file_path,
        "model": MODEL_NAME,
        "rules_loaded": len(rules_list),
        "total_runs_scanned": total_runs_scanned,
        "protected_bold_runs": protected_runs,
        "editable_runs": len(jobs),
        "rewritten_runs": rewritten_runs,
        "failed_runs": failed_runs
    }
    return stats


# ===============================================================
# 9) RUN
# ===============================================================

# Notebook usage:
stats = await transform_word_document(
    input_path=pathStr,
    output_file_path=output_path,
    category=document_category,
    rules_path=rules_file_path
)

print("Done. Summary:")
print(stats)

# Script usage (outside notebook):
# if __name__ == "__main__":
#     result = asyncio.run(transform_word_document(pathStr, output_path, document_category, rules_file_path))
#     print(result)

